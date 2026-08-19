"""Optional built-in tools for reading and writing PDF and DOCX files."""
from __future__ import annotations

import html
import os
import re
import stat
import tempfile
from pathlib import Path
from typing import Any, Callable

from .cancellation import CancellationToken, cancellation_or_default
from .models import ToolSpec
from .tools import Tool, ToolEffect, bound_output


def _path(arguments: dict[str, Any], extension: str) -> Path:
    raw_path = arguments.get("path")
    if not isinstance(raw_path, str) or not raw_path:
        raise ValueError("path is required")
    path = Path(raw_path)
    if path.suffix.casefold() != extension:
        raise ValueError(f"path must end with {extension}")
    return path


def _content(arguments: dict[str, Any]) -> str:
    content = arguments.get("content")
    if not isinstance(content, str):
        raise ValueError("content must be a string")
    return content


def _missing_dependency(tool: str, package: str) -> RuntimeError:
    return RuntimeError(
        f"{tool} requires {package}; install HAL document support with "
        'python -m pip install -e ".[documents]"'
    )


def _atomic_document_write(
    path: Path,
    writer: Callable[[Path], None],
    cancellation: CancellationToken | None,
) -> int:
    cancellation = cancellation_or_default(cancellation)
    cancellation.raise_if_cancelled()
    path.parent.mkdir(parents=True, exist_ok=True)
    mode = stat.S_IMODE(path.stat().st_mode) if path.exists() else 0o644
    descriptor, temp_name = tempfile.mkstemp(
        prefix=f".{path.stem}.", suffix=path.suffix, dir=path.parent,
    )
    os.close(descriptor)
    try:
        temporary = Path(temp_name)
        writer(temporary)
        cancellation.raise_if_cancelled()
        os.chmod(temporary, mode)
        size = temporary.stat().st_size
        os.replace(temporary, path)
        return size
    finally:
        if os.path.exists(temp_name):
            os.unlink(temp_name)


class PdfReadTool(Tool):
    parallel_safe = True
    effect = ToolEffect.READ_ONLY

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            "pdf_read",
            "Extract text from a PDF. start_page is a positive 1-indexed page; page_limit limits the number of pages.",
            {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "start_page": {"type": "integer", "minimum": 1},
                    "page_limit": {"type": "integer", "minimum": 1},
                },
                "required": ["path"],
            },
        )

    def run(
        self, arguments: dict[str, Any],
        cancellation: CancellationToken | None = None,
    ) -> str:
        cancellation = cancellation_or_default(cancellation)
        cancellation.raise_if_cancelled()
        path = _path(arguments, ".pdf")
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise _missing_dependency("pdf_read", "pypdf") from exc
        start_page = int(arguments.get("start_page", 1))
        raw_limit = arguments.get("page_limit")
        page_limit = int(raw_limit) if raw_limit is not None else None
        if start_page < 1:
            raise ValueError("start_page must be positive")
        if page_limit is not None and page_limit < 1:
            raise ValueError("page_limit must be positive")
        reader = PdfReader(str(path))
        total = len(reader.pages)
        if start_page > total and total:
            raise ValueError(f"start_page {start_page} is past the final page ({total})")
        stop = total if page_limit is None else min(total, start_page - 1 + page_limit)
        pages = []
        for index in range(start_page - 1, stop):
            cancellation.raise_if_cancelled()
            text = reader.pages[index].extract_text() or ""
            pages.append(f"--- Page {index + 1} of {total} ---\n{text.strip()}")
        return bound_output("\n\n".join(pages))


class PdfWriteTool(Tool):
    effect = ToolEffect.MUTATING

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            "pdf_write",
            "Create or replace a PDF from plain text. Markdown-style # headings are rendered as headings.",
            {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"},
                    "title": {"type": "string"},
                    "page_size": {"type": "string", "enum": ["letter", "a4"]},
                },
                "required": ["path", "content"],
            },
        )

    def run(
        self, arguments: dict[str, Any],
        cancellation: CancellationToken | None = None,
    ) -> str:
        path = _path(arguments, ".pdf")
        content = _content(arguments)
        title = arguments.get("title", "")
        if not isinstance(title, str):
            raise ValueError("title must be a string")
        page_size_name = str(arguments.get("page_size", "letter")).lower()
        if page_size_name not in {"letter", "a4"}:
            raise ValueError("page_size must be letter or a4")
        try:
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise _missing_dependency("pdf_write", "reportlab") from exc

        def write(temporary: Path) -> None:
            styles = getSampleStyleSheet()
            document = SimpleDocTemplate(
                str(temporary), pagesize=A4 if page_size_name == "a4" else letter,
                title=title or None, leftMargin=.75 * inch, rightMargin=.75 * inch,
                topMargin=.75 * inch, bottomMargin=.75 * inch,
            )
            story = []
            for block in re.split(r"\n\s*\n", content.strip()):
                if not block:
                    continue
                match = re.match(r"^(#{1,3})\s+(.+)$", block, re.DOTALL)
                if match:
                    style = styles[f"Heading{len(match.group(1))}"]
                    text = match.group(2)
                else:
                    style = styles["BodyText"]
                    text = block
                story.append(Paragraph(html.escape(text).replace("\n", "<br/>"), style))
                story.append(Spacer(1, 8))
            if not story:
                story.append(Paragraph("", styles["BodyText"]))
            document.build(story)

        size = _atomic_document_write(path, write, cancellation)
        return f"wrote {size} bytes to {path}"


class PdfFormWriteTool(Tool):
    effect = ToolEffect.MUTATING

    @property
    def spec(self) -> ToolSpec:
        position = {
            "x": {"type": "number", "minimum": 0},
            "y": {"type": "number", "minimum": 0},
        }
        return ToolSpec(
            "pdf_form_write",
            "Create a coordinate-controlled fillable AcroForm PDF. Coordinates are PDF points from the bottom-left. Signature fields are blank fields for later signing in Acrobat, not cryptographic signatures.",
            {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "title": {"type": "string"},
                    "page_size": {"type": "string", "enum": ["letter", "a4"]},
                    "pages": {
                        "type": "array", "minItems": 1,
                        "items": {
                            "type": "object",
                            "properties": {
                                "texts": {
                                    "type": "array", "items": {
                                        "type": "object",
                                        "properties": {
                                            "text": {"type": "string"}, **position,
                                            "font_size": {"type": "number", "minimum": 1},
                                            "bold": {"type": "boolean"},
                                        },
                                        "required": ["text", "x", "y"],
                                    },
                                },
                                "fields": {
                                    "type": "array", "items": {
                                        "type": "object",
                                        "properties": {
                                            "name": {"type": "string"},
                                            "type": {"type": "string", "enum": [
                                                "text", "multiline", "checkbox",
                                                "choice", "signature",
                                            ]},
                                            "label": {"type": "string"}, **position,
                                            "width": {"type": "number", "minimum": 1},
                                            "height": {"type": "number", "minimum": 1},
                                            "value": {},
                                            "options": {"type": "array", "items": {"type": "string"}},
                                            "required": {"type": "boolean"},
                                            "font_size": {"type": "number", "minimum": 1},
                                        },
                                        "required": ["name", "type", "x", "y", "width", "height"],
                                    },
                                },
                            },
                        },
                    },
                },
                "required": ["path", "pages"],
            },
        )

    def run(
        self, arguments: dict[str, Any],
        cancellation: CancellationToken | None = None,
    ) -> str:
        path = _path(arguments, ".pdf")
        title = arguments.get("title", "")
        pages = arguments.get("pages")
        page_size_name = str(arguments.get("page_size", "letter")).lower()
        if not isinstance(title, str):
            raise ValueError("title must be a string")
        if not isinstance(pages, list) or not pages:
            raise ValueError("pages must be a non-empty list")
        if page_size_name not in {"letter", "a4"}:
            raise ValueError("page_size must be letter or a4")
        try:
            from pypdf import PdfReader, PdfWriter
            from pypdf.generic import ArrayObject, NameObject, NumberObject
            from reportlab.lib.colors import black, white
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.pdfgen.canvas import Canvas
        except ImportError as exc:
            raise _missing_dependency("pdf_form_write", "pypdf and reportlab") from exc

        page_size = A4 if page_size_name == "a4" else letter
        page_width, page_height = page_size
        signatures: list[tuple[int, dict[str, Any]]] = []
        names: set[str] = set()

        def number(value: Any, label: str, *, positive: bool = False) -> float:
            if isinstance(value, bool) or not isinstance(value, (int, float)):
                raise ValueError(f"{label} must be a number")
            result = float(value)
            if result < (1 if positive else 0):
                raise ValueError(f"{label} must be {'positive' if positive else 'non-negative'}")
            return result

        def write(temporary: Path) -> None:
            canvas = Canvas(str(temporary), pagesize=page_size, pageCompression=1)
            if title:
                canvas.setTitle(title)
            for page_index, page in enumerate(pages):
                cancellation_or_default(cancellation).raise_if_cancelled()
                if not isinstance(page, dict):
                    raise ValueError(f"pages[{page_index}] must be an object")
                texts = page.get("texts", [])
                fields = page.get("fields", [])
                if not isinstance(texts, list) or not isinstance(fields, list):
                    raise ValueError(f"pages[{page_index}].texts and fields must be lists")
                for text_index, item in enumerate(texts):
                    if not isinstance(item, dict) or not isinstance(item.get("text"), str):
                        raise ValueError(f"pages[{page_index}].texts[{text_index}] needs string text")
                    x = number(item.get("x"), "text x")
                    y = number(item.get("y"), "text y")
                    font_size = number(item.get("font_size", 11), "font_size", positive=True)
                    if x > page_width or y > page_height:
                        raise ValueError("text coordinates are outside the page")
                    canvas.setFont("Helvetica-Bold" if item.get("bold") is True else "Helvetica", font_size)
                    canvas.drawString(x, y, item["text"])
                for field_index, field in enumerate(fields):
                    if not isinstance(field, dict):
                        raise ValueError(f"pages[{page_index}].fields[{field_index}] must be an object")
                    name = field.get("name")
                    kind = field.get("type")
                    if not isinstance(name, str) or not name.strip():
                        raise ValueError("form field name must be a non-empty string")
                    if name in names:
                        raise ValueError(f"duplicate form field name: {name}")
                    names.add(name)
                    if kind not in {"text", "multiline", "checkbox", "choice", "signature"}:
                        raise ValueError(f"unsupported form field type: {kind}")
                    x = number(field.get("x"), f"field {name} x")
                    y = number(field.get("y"), f"field {name} y")
                    width = number(field.get("width"), f"field {name} width", positive=True)
                    height = number(field.get("height"), f"field {name} height", positive=True)
                    if x + width > page_width or y + height > page_height:
                        raise ValueError(f"form field {name} extends outside the page")
                    label = field.get("label", "")
                    if not isinstance(label, str):
                        raise ValueError(f"form field {name} label must be a string")
                    if label:
                        canvas.setFont("Helvetica", 9)
                        canvas.drawString(x, min(page_height - 9, y + height + 3), label)
                    flags = "required" if field.get("required") is True else ""
                    common = dict(
                        name=name, tooltip=label or name, x=x, y=y,
                        borderColor=black, fillColor=white, borderWidth=1,
                        forceBorder=True,
                    )
                    if kind in {"text", "multiline"}:
                        value = field.get("value", "")
                        if not isinstance(value, str):
                            raise ValueError(f"form field {name} value must be a string")
                        field_flags = " ".join(filter(None, [flags, "multiline" if kind == "multiline" else ""]))
                        canvas.acroForm.textfield(
                            **common, width=width, height=height, value=value,
                            fieldFlags=field_flags,
                            fontSize=number(field.get("font_size", 10), "font_size", positive=True),
                        )
                    elif kind == "checkbox":
                        value = field.get("value", False)
                        if not isinstance(value, bool):
                            raise ValueError(f"form field {name} value must be true or false")
                        canvas.acroForm.checkbox(
                            **common, size=min(width, height), checked=value,
                            fieldFlags=flags,
                        )
                    elif kind == "choice":
                        options = field.get("options")
                        if not isinstance(options, list) or not options or any(not isinstance(option, str) for option in options):
                            raise ValueError(f"choice field {name} needs non-empty string options")
                        value = field.get("value", options[0])
                        if value not in options:
                            raise ValueError(f"choice field {name} value must be one of its options")
                        canvas.acroForm.choice(
                            **common, width=width, height=height, value=value,
                            options=options, fieldFlags=" ".join(filter(None, [flags, "combo"])),
                            fontSize=number(field.get("font_size", 10), "font_size", positive=True),
                        )
                    else:
                        canvas.rect(x, y, width, height, stroke=1, fill=0)
                        canvas.setFont("Helvetica-Oblique", 8)
                        canvas.drawCentredString(x + width / 2, y + height / 2 - 3, "Sign in Adobe Acrobat")
                        signatures.append((page_index, {
                            "name": name, "label": label or name,
                            "rect": [x, y, x + width, y + height],
                            "required": field.get("required") is True,
                        }))
                canvas.showPage()
            canvas.save()

            if signatures:
                reader = PdfReader(str(temporary))
                writer = PdfWriter()
                writer.clone_document_from_reader(reader)
                acroform = writer.root_object["/AcroForm"].get_object()
                fields = acroform.setdefault(NameObject("/Fields"), ArrayObject())
                acroform[NameObject("/SigFlags")] = NumberObject(3)
                for page_index, signature in signatures:
                    annotation = writer.add_annotation(page_index, {
                        "/Type": "/Annot", "/Subtype": "/Widget", "/FT": "/Sig",
                        "/T": signature["name"], "/TU": signature["label"],
                        "/Rect": signature["rect"], "/F": 4,
                        "/Ff": 2 if signature["required"] else 0,
                    })
                    fields.append(annotation.indirect_reference)
                second = temporary.with_name(f"{temporary.stem}.fields{temporary.suffix}")
                try:
                    with second.open("wb") as stream:
                        writer.write(stream)
                    os.replace(second, temporary)
                finally:
                    if second.exists():
                        second.unlink()

        size = _atomic_document_write(path, write, cancellation)
        return f"wrote fillable form with {len(names)} fields ({len(signatures)} signature) and {len(pages)} pages; {size} bytes to {path}"


class DocxReadTool(Tool):
    parallel_safe = True
    effect = ToolEffect.READ_ONLY

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            "docx_read",
            "Extract paragraphs and tables from a DOCX document as plain text.",
            {
                "type": "object",
                "properties": {"path": {"type": "string"}},
                "required": ["path"],
            },
        )

    def run(
        self, arguments: dict[str, Any],
        cancellation: CancellationToken | None = None,
    ) -> str:
        cancellation = cancellation_or_default(cancellation)
        cancellation.raise_if_cancelled()
        path = _path(arguments, ".docx")
        try:
            from docx import Document
        except ImportError as exc:
            raise _missing_dependency("docx_read", "python-docx") from exc
        document = Document(str(path))
        output = []
        for paragraph in document.paragraphs:
            cancellation.raise_if_cancelled()
            text = paragraph.text.strip()
            if text:
                style = paragraph.style.name if paragraph.style is not None else ""
                match = re.fullmatch(r"Heading ([1-6])", style)
                output.append(f"{'#' * int(match.group(1))} {text}" if match else text)
        for number, table in enumerate(document.tables, 1):
            cancellation.raise_if_cancelled()
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            output.append(f"[Table {number}]\n" + "\n".join(rows))
        return bound_output("\n\n".join(output))


class DocxWriteTool(Tool):
    effect = ToolEffect.MUTATING

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            "docx_write",
            "Create or replace a DOCX from plain text. Markdown-style # headings and simple lists are rendered structurally.",
            {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"},
                    "title": {"type": "string"},
                },
                "required": ["path", "content"],
            },
        )

    def run(
        self, arguments: dict[str, Any],
        cancellation: CancellationToken | None = None,
    ) -> str:
        path = _path(arguments, ".docx")
        content = _content(arguments)
        title = arguments.get("title", "")
        if not isinstance(title, str):
            raise ValueError("title must be a string")
        try:
            from docx import Document
        except ImportError as exc:
            raise _missing_dependency("docx_write", "python-docx") from exc

        def write(temporary: Path) -> None:
            document = Document()
            if title:
                document.core_properties.title = title
            for block in re.split(r"\n\s*\n", content.strip()):
                if not block:
                    continue
                heading = re.match(r"^(#{1,6})\s+(.+)$", block, re.DOTALL)
                if heading:
                    document.add_heading(heading.group(2), level=len(heading.group(1)))
                    continue
                lines = block.splitlines()
                if all(line.startswith("- ") for line in lines):
                    for line in lines:
                        document.add_paragraph(line[2:], style="List Bullet")
                elif all(re.match(r"\d+[.)]\s+", line) for line in lines):
                    for line in lines:
                        document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
                else:
                    document.add_paragraph("\n".join(lines))
            document.save(str(temporary))

        size = _atomic_document_write(path, write, cancellation)
        return f"wrote {size} bytes to {path}"
